from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run('Base Layer')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
title.paragraph_format.space_after = Pt(16)

def add_heading(text):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x2a, 0x2a, 0x2a)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)

def add_para(text):
    p = doc.add_paragraph()
    parts = text.split('**')
    for i, part in enumerate(parts):
        if part:
            run = p.add_run(part)
            run.font.size = Pt(11)
            if i % 2 == 1:
                run.bold = True
    p.paragraph_format.space_after = Pt(6)

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    parts = text.split('**')
    for i, part in enumerate(parts):
        if part:
            run = p.add_run(part)
            run.font.size = Pt(11)
            if i % 2 == 1:
                run.bold = True
    p.paragraph_format.space_after = Pt(4)

# === THE PROBLEM ===
add_heading('The Problem')

add_para('Every major AI platform now offers some form of memory. ChatGPT remembers preferences. Claude retains project context. But none of it is yours. Your understanding of yourself, built across thousands of conversations over years, is siloed inside each platform and owned by the provider. There is no universal memory layer. There is no original timeline. Switch providers, and you start from zero.')

add_para('Platform memory is also shallow; it stores what you said, not how you think. Dumping conversation history into a prompt is not memory; it is a filing cabinet with no librarian. The problem is not that AI lacks recall. The problem is that no one is curating what actually matters about a person, and that curation must be owned by the individual, not the platform.')

add_para('Other companies are working in this space. Limitless records and retrieves personal context, Mem and Personal.ai offer persistent memory features, Kin is building long-term personal AI. Most of these are building better retrieval: smarter filing cabinets. Base Layer is building behavioral understanding. Not what you said, but how you think, decide, and respond. That distinction is the difference between an AI that recalls and one that knows you.')

# === THE SOLUTION ===
add_heading('The Solution')

add_para('Base Layer is a portable memory system owned by the individual, not the platform. It curates what matters about a person from their unstructured data and compresses it into a behavioral brief that any AI can use. The system models how you think, decide, and respond. The brief travels with you across providers.')

add_para('The core principle: you cannot accurately model a complex system without understanding its composite parts. Before you can model a team, an organization, or a system of people, you need to model one person well. That is the base layer.')

# === WHY ME ===
add_heading('Why Me')

add_para('I have spent the last four years at the intersection of AI, context management, and complex systems. At a previous startup (accelerator-backed), I built and ran an operations platform for safety-critical software systems, working with enterprise and government clients. The core problem was the same one I am solving now: curating exactly the right context from vast distributed data so an intelligence can reason accurately. I raised pre-seed funding, closed enterprise contracts, and led GTM, operations, and customer programs across federal and commercial verticals for 3.5 years.')

add_para('The company shut down. The problem stayed with me. What that experience taught me is that the people who manage context best do not hoard information. They curate it. But we were modeling systems, not the people operating them. The individual was the missing layer.')

# === HOW IT WORKS ===
add_heading('How It Works')

add_para('The system is a five-layer pipeline. Raw conversations are stored locally in SQLite as immutable ground truth. Content is embedded as vectors for semantic retrieval. A local 14B-parameter model extracts atomic facts, scores them for novelty and significance, deduplicates against existing knowledge, and manages a full lifecycle (add, update, supersede, skip) for every piece of information. Corrections propagate permanently. Retrieved memory is compressed into a three-tier brief: an always-on identity block (behavioral predictions about who you are), a dynamic theme block (facts relevant to the current conversation), and an episodic block (recent interaction context). The assembled brief is injected into any LLM that accepts a system prompt. All personal data stays on the user\u2019s device. Only the compressed brief reaches the cloud.')

# === WHAT HAS BEEN BUILT ===
add_heading('What Has Been Built')

add_para('The system has been built over 24 working sessions, each solving a specific failure mode of the previous.')

add_bullet('**4,700+ extracted facts** from 1,800+ conversations spanning 3 years of AI use')
add_bullet('**10 universal identity clusters** organizing facts into behavioral dimensions (who you are, what you\u2019ve built, what drives you, how you operate, etc.)')
add_bullet('**Evaluation infrastructure:** automated test harness (85% pass rate, zero hallucinated data leakage), self-correcting knowledge base where corrections propagate permanently')
add_bullet('**Full pipeline running on consumer hardware** (NVIDIA RTX 3080, 10GB VRAM)')
add_bullet('**37 documented design decisions** with full reasoning')

# === WHERE THIS GOES ===
add_heading('Where This Goes')

add_para('**Individual, then multi-user, then enterprise.** The immediate priority is multi-user validation: proving the system generalizes beyond the builder. From there, a public prototype where anyone can connect their data and experience persistent AI understanding firsthand.')

add_para('**Revenue model:** Open-core. The self-hosted pipeline is free for power users who run it locally. A hosted tier serves non-technical users who want persistent memory without managing infrastructure. Enterprise licensing for organizational memory. The core IP (extraction methodology, clustering logic, scoring algorithms, identity block authoring) is the moat.')

add_para('**The long-term opportunity is portable memory.** Today, when someone leaves a company, their context leaves with them and institutional knowledge is lost overnight. Base Layer inverts this. Individuals own their memory layer. When they join an organization, their model integrates. When they leave, they take their memory with them. The organization keeps what was shared; the individual keeps what is theirs. Separating personal from institutional knowledge is a hard problem, technically, legally, and organizationally. Base Layer starts with the individual because that is the foundational unit, and the one where the approach can be validated before the complexity scales. Portable memory is also the prerequisite for what follows: agents that act on your behalf because they actually know you, intelligence that persists as models upgrade underneath, and collective understanding that emerges without erasing the individual.')

add_para('**The work ahead is not only technical.** The AI memory space is moving fast, and building the right system is necessary but not sufficient. Staying ahead means tracking how platforms evolve their memory capabilities, where gaps remain for a standalone layer, and how data ownership norms shift. Personal memory is deeply sensitive data, and the business model must respect that. Local-first architecture is not just a technical choice; it is an ethical one. The system that knows you best should be the one you control. Building something both profitable and trustworthy is the harder path, but it is the only one that produces a durable company in this space.')

from pathlib import Path
output_path = str(Path(__file__).parent.parent / "gtm" / "generated" / "EWOR_APPLICATION.docx")
doc.save(output_path)
print(f'Saved to {output_path}')
